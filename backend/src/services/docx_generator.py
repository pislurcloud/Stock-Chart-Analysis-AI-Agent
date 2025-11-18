"""
DOCX Report Generator
Converts markdown analysis reports to professional Word documents
"""

import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXReportGenerator:
    """
    Generates professional DOCX reports from analysis data
    """
    
    def __init__(self):
        """Initialize the DOCX generator"""
        self.output_dir = os.getenv("DOCX_OUTPUT_DIR", "./reports")
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_report(
        self,
        analysis_result: Dict[str, Any],
        include_chart: bool = True
    ) -> str:
        """
        Generate professional DOCX report
        
        Args:
            analysis_result: Complete analysis from AIAnalysisCrew
            include_chart: Whether to embed the chart image
        
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating DOCX report for {analysis_result['symbol']}")
        
        try:
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_styles(doc)
            
            # Add title page
            self._add_title_page(doc, analysis_result)
            
            # Add executive summary
            self._add_executive_summary(doc, analysis_result)
            
            # Add technical analysis section
            self._add_technical_analysis(doc, analysis_result)
            
            # Add scenario playbook
            self._add_scenario_playbook(doc, analysis_result)
            
            # Add pattern analysis
            self._add_pattern_analysis(doc, analysis_result)
            
            # Add risk metrics
            self._add_risk_metrics(doc, analysis_result)
            
            # Add chart if available
            if include_chart and analysis_result.get('chart_path'):
                self._add_chart(doc, analysis_result)
            
            # Add disclaimer
            self._add_disclaimer(doc)
            
            # Save document
            filename = self._generate_filename(analysis_result)
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            logger.info(f"DOCX report saved: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            raise
    
    def _setup_styles(self, doc: Document):
        """Set up custom styles for the document"""
        
        # Heading 1 style
        if 'Custom Heading 1' not in doc.styles:
            h1_style = doc.styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.size = Pt(18)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    def _add_title_page(self, doc: Document, data: Dict[str, Any]):
        """Add professional title page"""
        
        # Company name and symbol
        title = doc.add_heading(f"{data['stock_info']['company_name']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"({data['symbol']} - {data['timeframe'].upper()})")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(14)
        subtitle_format.font.color.rgb = RGBColor(64, 64, 64)
        
        # Recommendation badge
        doc.add_paragraph()
        recommendation = doc.add_paragraph(f"📊 {data['report']['recommendation']}")
        recommendation.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rec_format = recommendation.runs[0]
        rec_format.font.size = Pt(16)
        rec_format.font.bold = True
        
        # Metadata
        doc.add_paragraph()
        doc.add_paragraph()
        
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        
        meta_data = [
            ('Current Price', f"₹{data['latest_candle']['close']:.2f}"),
            ('Volume', f"{data['latest_candle']['volume']:,}"),
            ('Sector', data['stock_info'].get('sector', 'N/A')),
            ('Analysis Date', datetime.now().strftime('%Y-%m-%d %H:%M IST')),
            ('Execution Time', f"{data['execution_time']:.2f}s")
        ]
        
        for i, (label, value) in enumerate(meta_data):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[1].text = value
            meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """Add executive summary section"""
        
        doc.add_heading('Executive Summary', 1)
        
        summary_text = data['report']['summary']
        doc.add_paragraph(summary_text)
        
        doc.add_paragraph()
        
        # Key metrics
        doc.add_heading('Key Metrics', 2)
        
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Light List Accent 1'
        
        metrics = [
            ('Technical Bias', data['technical_analysis']['overall_bias']),
            ('Signal Strength', f"{data['technical_analysis']['strength_score']}/100"),
            ('Trend', data['technical_analysis']['trend']),
            ('Market Structure', data['pattern_analysis']['market_structure'])
        ]
        
        for i, (label, value) in enumerate(metrics):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = str(value)
            metrics_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_page_break()
    
    def _add_technical_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add technical analysis details"""
        
        doc.add_heading('Technical Analysis', 1)
        
        ta = data['technical_analysis']
        
        # Overall assessment
        doc.add_heading('Overall Assessment', 2)
        doc.add_paragraph(f"The technical analysis shows a {ta['overall_bias']} bias with a signal strength of {ta['strength_score']}/100.")
        
        doc.add_paragraph()
        
        # Key levels
        doc.add_heading('Key Price Levels', 2)
        
        levels_table = doc.add_table(rows=len(ta['key_levels']) + 1, cols=2)
        levels_table.style = 'Medium Shading 1 Accent 1'
        
        # Header
        levels_table.rows[0].cells[0].text = 'Level'
        levels_table.rows[0].cells[1].text = 'Price (₹)'
        
        for i, cell in enumerate(levels_table.rows[0].cells):
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data
        for i, (level, price) in enumerate(ta['key_levels'].items(), start=1):
            levels_table.rows[i].cells[0].text = level.replace('_', ' ').title()
            levels_table.rows[i].cells[1].text = f"₹{price:.2f}"
        
        doc.add_page_break()
    
    def _add_scenario_playbook(self, doc: Document, data: Dict[str, Any]):
        """Add scenario playbook section with formatted table"""
        
        doc.add_heading('🧭 Scenario Playbook', 1)
        
        doc.add_paragraph("Trade scenarios with entry, stop loss, and target levels:")
        doc.add_paragraph()
        
        scenarios = data['risk_analysis']['scenarios']
        
        if scenarios:
            # Create table
            table = doc.add_table(rows=len(scenarios) + 1, cols=7)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            headers = ['Scenario', 'Entry', 'Stop', 'Target', 'R:R', 'Confidence', 'WarrenAI Take']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Data rows
            for i, scenario in enumerate(scenarios, start=1):
                row = table.rows[i]
                row.cells[0].text = scenario['scenario']
                row.cells[1].text = f"₹{scenario['entry']:.2f}"
                row.cells[2].text = f"₹{scenario['stop']:.2f}"
                row.cells[3].text = f"₹{scenario['target']:.2f}"
                row.cells[4].text = str(scenario['rr_ratio'])
                row.cells[5].text = scenario['confidence']
                row.cells[6].text = scenario['warreni_take']
                
                # Font size for readability
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            doc.add_paragraph()
            
            # Insights
            doc.add_heading('💡 Insights', 2)
            
            position_sizing = data['risk_analysis']['position_sizing']
            doc.add_paragraph(f"• Recommended Position Size: {position_sizing['recommended_position']}")
            doc.add_paragraph(f"• Rationale: {position_sizing['rationale']}")
            doc.add_paragraph(f"• Max Risk per Trade: {position_sizing['max_risk_per_trade']}")
        
        doc.add_page_break()
    
    def _add_pattern_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add pattern analysis section"""
        
        doc.add_heading('🧩 Pattern Analysis', 1)
        
        pattern_data = data['pattern_analysis']
        
        doc.add_paragraph(f"Market Structure: {pattern_data['market_structure']}")
        doc.add_paragraph(f"Confidence: {pattern_data['confidence']}")
        
        doc.add_paragraph()
        
        # Patterns identified
        if pattern_data.get('patterns'):
            doc.add_heading('Patterns Identified', 2)
            for pattern in pattern_data['patterns'][:5]:
                doc.add_paragraph(f"• {pattern.get('pattern', 'Unknown pattern')}", style='List Bullet')
        
        doc.add_page_break()
    
    def _add_risk_metrics(self, doc: Document, data: Dict[str, Any]):
        """Add risk metrics section"""
        
        doc.add_heading('⚖️ Risk Metrics', 1)
        
        risk_metrics = data['risk_analysis']['risk_metrics']
        
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Medium Shading 1 Accent 1'
        
        metrics = [
            ('Average R:R Ratio', str(risk_metrics['average_rr_ratio'])),
            ('Volatility Level', risk_metrics['volatility_level']),
            ('Risk Grade', risk_metrics['risk_grade']),
            ('Expected Range', risk_metrics['expected_range'])
        ]
        
        for i, (label, value) in enumerate(metrics):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = value
            metrics_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Backtest results
        doc.add_heading('📈 Backtest Results', 2)
        backtest = data['risk_analysis']['backtest']
        doc.add_paragraph(f"Estimated Success Rate: {backtest['estimated_success_rate']}%")
        doc.add_paragraph(f"Lookback Period: {backtest['lookback_period']}")
        doc.add_paragraph(f"Note: {backtest['note']}")
        
        doc.add_page_break()
    
    def _add_chart(self, doc: Document, data: Dict[str, Any]):
        """Add chart image to document"""
        
        doc.add_heading('📊 Technical Chart', 1)
        
        chart_path = data['chart_path']
        
        if os.path.exists(chart_path):
            try:
                doc.add_picture(chart_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.info("Chart embedded successfully")
            except Exception as e:
                logger.warning(f"Could not embed chart: {str(e)}")
                doc.add_paragraph(f"Chart available at: {chart_path}")
        else:
            doc.add_paragraph("Chart not available")
        
        doc.add_page_break()
    
    def _add_disclaimer(self, doc: Document):
        """Add disclaimer section"""
        
        doc.add_heading('⚠️ Disclaimer', 1)
        
        disclaimer_text = """
This analysis is for educational and informational purposes only. It is NOT financial advice, 
investment advice, or a recommendation to buy or sell any securities.

All trading and investment decisions should be made based on your own research, risk tolerance, 
and financial situation. Past performance does not guarantee future results.

Technical analysis and AI-generated insights are tools to assist in decision-making but should 
not be relied upon as the sole basis for trading decisions.

Always consult with a qualified financial advisor before making investment decisions.

The creators of this analysis are not responsible for any losses incurred from trading decisions 
made based on this information.
        """
        
        disclaimer_para = doc.add_paragraph(disclaimer_text)
        disclaimer_para.runs[0].font.size = Pt(9)
        disclaimer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S IST')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    def _generate_filename(self, data: Dict[str, Any]) -> str:
        """Generate filename for the DOCX report"""
        
        symbol = data['symbol']
        timeframe = data['timeframe']
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"{symbol}_{timeframe}_analysis_{timestamp}.docx"


# Test function
if __name__ == "__main__":
    print("="*80)
    print("Testing DOCX Report Generator")
    print("="*80)
    
    # This would normally receive data from the AI crew
    # For testing, we'd need actual analysis data
    print("\nDOCX generator ready!")
    print("Use generate_report() with analysis results from AIAnalysisCrew")
    print("="*80)